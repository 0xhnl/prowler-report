import argparse
import re
from pathlib import Path

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HTML_COLUMNS_TO_EXTRACT = [0, 1, 2, 3, 5, 6, 8, 9, 10]
CSV_HEADER = [
    "Status",
    "Severity",
    "Service Name",
    "Region",
    "Check Title",
    "Resource ID",
    "Status Extended",
    "Risk",
    "Recommendation",
]

REMEDIATION_CODE_FIELDS = [
    ("CLI", "REMEDIATION_CODE_CLI"),
    ("NativeIaC", "REMEDIATION_CODE_NATIVEIAC"),
    ("Terraform", "REMEDIATION_CODE_TERRAFORM"),
]

INLINE_MARKDOWN_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`\n]+`)")
FENCED_CODE_PATTERN = re.compile(r"```([^\n`]*)\n?(.*?)```", re.DOTALL)
CODE_FONT = "Consolas"
CODE_BOX_FILL = "F2F2F2"


def _set_shading(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _add_inline_runs(paragraph, text):
    """Add runs to paragraph, converting **bold** and `code` markdown."""
    pos = 0
    for match in INLINE_MARKDOWN_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_rich_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_inline_runs(paragraph, text)
    return paragraph


def _add_rich_bullet(doc, text):
    """Render text as List Bullets, splitting paragraphs on blank lines and
    promoting '- ' lines to List Bullet 2 sub-bullets."""
    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = paragraph.strip("\n")
        if not paragraph.strip():
            continue
        lines = paragraph.split("\n")
        buffer = []
        first_text_block = True

        def flush():
            nonlocal first_text_block
            if not buffer:
                return
            merged = "\n".join(buffer).strip()
            buffer.clear()
            if not merged:
                return
            style = "List Bullet" if first_text_block else None
            first_text_block = False
            _add_rich_paragraph(doc, merged, style=style)

        for line in lines:
            stripped = line.lstrip()
            if stripped.startswith("- "):
                flush()
                _add_rich_paragraph(doc, stripped[2:].strip(), style="List Bullet 2")
            elif re.match(r"^\d+\.\s+", stripped):
                flush()
                _add_rich_paragraph(doc, stripped, style="List Bullet")
                first_text_block = False
            else:
                buffer.append(line)
        flush()


def _add_code_box(doc, code, lang=None):
    """Render code inside a shaded paragraph with a monospace font."""
    paragraph = doc.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()
    _set_shading(p_pr, CODE_BOX_FILL)
    if lang:
        header = paragraph.add_run(f"{lang}\n")
        header.bold = True
        header.font.name = CODE_FONT
    body = paragraph.add_run(code)
    body.font.name = CODE_FONT


def _add_rich_block(doc, text):
    """Render text that may include fenced ```code``` blocks as code boxes."""
    pos = 0
    for match in FENCED_CODE_PATTERN.finditer(text):
        prefix = text[pos:match.start()].strip()
        if prefix:
            _add_rich_paragraph(doc, prefix)
        lang = match.group(1).strip() or None
        code = match.group(2).strip("\n")
        _add_code_box(doc, code, lang=lang)
        pos = match.end()
    trailing = text[pos:].strip()
    if trailing:
        _add_rich_paragraph(doc, trailing)


def clean_text(text):
    """Remove HTML artifacts and normalize whitespace."""
    text = text.replace("<wbr />", "")
    soup = BeautifulSoup(text, "html.parser")
    text = soup.get_text(separator=" ", strip=True)
    return re.sub(r"\s+", " ", text).strip()


def parse_prowler_html(input_file):
    """Parse the Prowler HTML report into a DataFrame."""
    html_content = Path(input_file).read_text(encoding="utf-8")
    soup = BeautifulSoup(html_content, "html.parser")

    findings_table = soup.find("table", id="findingsTable")
    if not findings_table:
        raise ValueError("Could not find the main findings table (id='findingsTable') in the HTML.")

    tbody = findings_table.find("tbody")
    if not tbody:
        raise ValueError("Could not find the findings table body in the HTML.")

    findings = []
    for row in tbody.find_all("tr"):
        cells = row.find_all("td")
        if len(cells) < 11:
            continue

        row_data = []
        for index in HTML_COLUMNS_TO_EXTRACT:
            row_data.append(clean_text(cells[index].decode_contents()))
        findings.append(row_data)

    df = pd.DataFrame(findings, columns=CSV_HEADER)
    df["Remediation"] = ""
    df["Remediation Steps"] = ""
    return df


def parse_prowler_csv(input_file):
    """Parse the Prowler semicolon-delimited CSV into a DataFrame."""
    df = pd.read_csv(input_file, sep=";", dtype=str, keep_default_na=False)

    def combine_remediation_code(row):
        parts = []
        for label, column in REMEDIATION_CODE_FIELDS:
            value = str(row.get(column, "")).strip()
            if not value:
                continue
            if value.startswith("```"):
                parts.append(f"**{label}**\n{value}")
            else:
                parts.append(f"**{label}**\n```\n{value}\n```")
        return "\n\n".join(parts)

    out = pd.DataFrame(
        {
            "Status": df.get("STATUS", ""),
            "Severity": df.get("SEVERITY", ""),
            "Service Name": df.get("SERVICE_NAME", ""),
            "Region": df.get("REGION", ""),
            "Check Title": df.get("CHECK_TITLE", ""),
            "Resource ID": df.get("RESOURCE_UID", ""),
            "Status Extended": df.get("STATUS_EXTENDED", ""),
            "Risk": df.get("RISK", ""),
            "Recommendation": df.get("REMEDIATION_RECOMMENDATION_TEXT", ""),
            "Remediation": df.apply(combine_remediation_code, axis=1),
            "Remediation Steps": df.get("REMEDIATION_CODE_OTHER", ""),
        }
    )
    return out


def normalize_column_name(name):
    return "".join(ch.lower() for ch in str(name).strip() if ch.isalnum())


def resolve_column(df, candidates):
    normalized = {normalize_column_name(col): col for col in df.columns}
    for candidate in candidates:
        resolved = normalized.get(normalize_column_name(candidate))
        if resolved:
            return resolved
    return None


def build_docx(input_df, output_docx, severity):
    df = input_df.copy()
    df.columns = [c.strip() for c in df.columns]

    col_map = {
        "Severity": ["Severity"],
        "Status": ["Status"],
        "Check Title": ["Check Title", "Title", "CheckTitle"],
        "Status Extended": ["Status Extended", "StatusExtended", "Description"],
        "Risk": ["Risk"],
        "Resource ID": [
            "Resource ID",
            "ResourceID",
            "Affected Resource IDs",
            "resource_id",
        ],
        "Recommendation": ["Recommendation", "Recommendations"],
    }

    optional_col_map = {
        "Remediation": ["Remediation", "RemediationCode"],
        "Remediation Steps": ["Remediation Steps", "RemediationSteps", "RemediationRecommendationText"],
    }

    resolved = {}
    for key, names in col_map.items():
        found = resolve_column(df, names)
        if not found:
            raise ValueError(f"Missing required column for {key}")
        resolved[key] = found

    for key, names in optional_col_map.items():
        resolved[key] = resolve_column(df, names)

    df_filtered = df[
        (df[resolved["Status"]].astype(str).str.lower() == "fail")
        & (df[resolved["Severity"]].astype(str).str.lower() == severity.lower())
    ]

    doc = Document()
    doc.add_heading(f"{severity.title()} Findings", level=1)

    if df_filtered.empty:
        doc.add_paragraph(f"No FAIL findings found for severity: {severity.upper()}.")
        doc.save(output_docx)
        return 0

    grouped = df_filtered.groupby(resolved["Check Title"], dropna=False)

    for check_title, group in grouped:
        doc.add_heading(str(check_title), level=3)

        def add_unique_bullets(title, column):
            doc.add_heading(title, level=4)
            unique_values = group[column].dropna().astype(str).unique()
            if len(unique_values) == 0:
                doc.add_paragraph("No data found.", style="List Bullet")
                return
            for val in unique_values:
                _add_rich_bullet(doc, val)

        def add_unique_rich_blocks(title, column):
            doc.add_heading(title, level=4)
            unique_values = group[column].dropna().astype(str).unique()
            if len(unique_values) == 0:
                doc.add_paragraph("No data found.")
                return
            for val in unique_values:
                _add_rich_block(doc, val)

        def add_remediation_section(title, text_column, code_column):
            doc.add_heading(title, level=4)
            rendered_any = False
            text_values = group[text_column].dropna().astype(str).unique() if text_column else []
            for val in text_values:
                if val.strip():
                    _add_rich_bullet(doc, val)
                    rendered_any = True
            code_values = group[code_column].dropna().astype(str).unique() if code_column else []
            for val in code_values:
                if val.strip():
                    _add_rich_block(doc, val)
                    rendered_any = True
            if not rendered_any:
                doc.add_paragraph("No data found.", style="List Bullet")

        add_unique_bullets("Description", resolved["Status Extended"])
        add_unique_bullets("Risk", resolved["Risk"])
        add_unique_bullets("Affected Resource IDs", resolved["Resource ID"])
        if resolved.get("Remediation"):
            add_remediation_section(
                "Remediation:",
                resolved["Recommendation"],
                resolved["Remediation"],
            )
        if resolved.get("Remediation Steps"):
            add_unique_bullets("Remediation Steps:", resolved["Remediation Steps"])
        doc.add_paragraph("")

    doc.save(output_docx)
    return len(df_filtered)


def main():
    parser = argparse.ArgumentParser(
        description="Parse a Prowler HTML report and generate severity-based DOCX reports."
    )
    parser.add_argument("-f", "--file", required=True, help="Input Prowler HTML or CSV file")
    parser.add_argument("-o", "--output", required=True, help="Output directory for DOCX files")
    args = parser.parse_args()

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    suffix = Path(args.file).suffix.lower()
    if suffix == ".csv":
        df = parse_prowler_csv(args.file)
    else:
        df = parse_prowler_html(args.file)

    severities = ["critical", "high", "medium", "low"]
    for severity in severities:
        output_docx = output_dir / f"{severity}.docx"
        count = build_docx(df, output_docx, severity)
        print(f"[+] Generated {output_docx} ({count} FAIL findings)")


if __name__ == "__main__":
    main()
